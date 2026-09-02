"""
TextLens — Document text extraction.
Supports PDF (PyMuPDF), DOCX (python-docx), TXT.
Returns a structured list of pages with raw text.
"""
from __future__ import annotations

import io
from dataclasses import dataclass, field
from pathlib import Path
from typing import List

from loguru import logger


@dataclass
class ExtractedPage:
    page_number: int  # 1-indexed
    text: str
    word_count: int = 0

    def __post_init__(self):
        self.word_count = len(self.text.split())


@dataclass
class ExtractionResult:
    pages: List[ExtractedPage] = field(default_factory=list)
    file_type: str = ""
    error: str = ""

    @property
    def success(self) -> bool:
        return not self.error

    @property
    def full_text(self) -> str:
        return "\n\n".join(p.text for p in self.pages)

    @property
    def page_count(self) -> int:
        return len(self.pages)


# ───────── PDF ─────────

def extract_pdf(path: Path) -> ExtractionResult:
    try:
        import fitz  # PyMuPDF
    except ImportError:
        return ExtractionResult(error="PyMuPDF not installed.")

    try:
        doc = fitz.open(str(path))
        pages: List[ExtractedPage] = []
        for i, page in enumerate(doc, start=1):
            text = page.get_text("text")
            if text.strip():
                pages.append(ExtractedPage(page_number=i, text=text.strip()))
        doc.close()

        if not pages:
            return ExtractionResult(error="No readable text found in PDF.")

        logger.info(f"PDF extracted: {len(pages)} pages from {path.name}")
        return ExtractionResult(pages=pages, file_type="pdf")

    except Exception as exc:
        logger.error(f"PDF extraction failed: {exc}")
        return ExtractionResult(error=f"PDF extraction error: {exc}")


# ───────── DOCX ─────────

def extract_docx(path: Path) -> ExtractionResult:
    try:
        from docx import Document as DocxDocument
    except ImportError:
        return ExtractionResult(error="python-docx not installed.")

    try:
        doc = DocxDocument(str(path))
        paragraphs: List[str] = []
        current_section = ""
        page_texts: List[str] = []
        current_page_paras: List[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Detect headings as section markers
            if para.style.name.startswith("Heading"):
                if current_page_paras:
                    page_texts.append("\n".join(current_page_paras))
                    current_page_paras = []
                current_page_paras.append(text)
            else:
                current_page_paras.append(text)

        if current_page_paras:
            page_texts.append("\n".join(current_page_paras))

        if not page_texts:
            # Fallback: treat whole doc as page 1
            all_text = "\n".join(
                p.text.strip() for p in doc.paragraphs if p.text.strip()
            )
            if not all_text:
                return ExtractionResult(error="DOCX contains no readable text.")
            page_texts = [all_text]

        pages = [
            ExtractedPage(page_number=i + 1, text=t)
            for i, t in enumerate(page_texts)
        ]
        logger.info(f"DOCX extracted: {len(pages)} sections from {path.name}")
        return ExtractionResult(pages=pages, file_type="docx")

    except Exception as exc:
        logger.error(f"DOCX extraction failed: {exc}")
        return ExtractionResult(error=f"DOCX extraction error: {exc}")


# ───────── TXT ─────────

def extract_txt(path: Path) -> ExtractionResult:
    try:
        import chardet

        raw = path.read_bytes()
        if not raw:
            return ExtractionResult(error="File is empty.")

        detected = chardet.detect(raw)
        encoding = detected.get("encoding") or "utf-8"

        try:
            text = raw.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            text = raw.decode("utf-8", errors="replace")

        text = text.strip()
        if not text:
            return ExtractionResult(error="TXT file contains no text after decoding.")

        # Split into pseudo-pages by blank lines (~500 words each)
        words = text.split()
        chunk_size = 500
        pages: List[ExtractedPage] = []
        for i in range(0, max(len(words), 1), chunk_size):
            chunk = " ".join(words[i : i + chunk_size])
            pages.append(ExtractedPage(page_number=len(pages) + 1, text=chunk))

        if not pages:
            pages = [ExtractedPage(page_number=1, text=text)]

        logger.info(f"TXT extracted: {len(pages)} chunks from {path.name}")
        return ExtractionResult(pages=pages, file_type="txt")

    except Exception as exc:
        logger.error(f"TXT extraction failed: {exc}")
        return ExtractionResult(error=f"TXT extraction error: {exc}")


# ───────── Dispatcher ─────────

EXTRACTORS = {
    ".pdf": extract_pdf,
    ".docx": extract_docx,
    ".txt": extract_txt,
}

ALLOWED_EXTENSIONS = set(EXTRACTORS.keys())


def extract_document(path: Path) -> ExtractionResult:
    """Entry point: detect type and extract."""
    suffix = path.suffix.lower()
    if suffix not in EXTRACTORS:
        return ExtractionResult(error=f"Unsupported file type: {suffix}")
    return EXTRACTORS[suffix](path)
